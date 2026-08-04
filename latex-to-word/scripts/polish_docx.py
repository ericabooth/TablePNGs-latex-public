#!/usr/bin/env python3
"""Post-conversion polish on a pandoc-produced .docx.

Fixes the things that are easier to correct in the finished document than in the
LaTeX source:

  * table cells default to bottom alignment, so when one header wraps to two
    lines the other headers drop and the row looks broken. Force top alignment.
  * header rows come through unbolded, which makes wide tables hard to scan.
  * header rows should repeat when a table splits across pages.

Usage: polish_docx.py IN.docx OUT.docx
"""
import sys, copy, docx
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

src, dst = sys.argv[1], sys.argv[2]
d = docx.Document(src)

tables = bolded = 0
for t in d.tables:
    tables += 1
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if not t.rows:
        continue
    hdr = t.rows[0]
    # bold the header text
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    bolded += 1
    # ask Word to repeat the header row on each page the table spans
    trPr = hdr._tr.get_or_add_trPr()
    if trPr.find(qn('w:tblHeader')) is None:
        el = OxmlElement('w:tblHeader')
        el.set(qn('w:val'), 'true')
        trPr.append(el)

d.save(dst)
print(f'polished {tables} tables (top-aligned cells, {bolded} bold repeating headers)')
