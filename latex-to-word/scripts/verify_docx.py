#!/usr/bin/env python3
"""Verification battery for a LaTeX-to-Word conversion.

Runs the checks that caught real defects on the reference jobs, so they run
every time instead of when someone remembers. Exit code is nonzero on failure.

Usage: verify_docx.py OUT.docx PREPPED.tex [ORIGINAL.pdf]

Requires: python-docx; pdftotext on PATH; LibreOffice (soffice) on PATH or at
the standard macOS location if OUT.pdf does not already exist beside OUT.docx.
"""
import re, sys, subprocess, pathlib, shutil
import docx
from collections import Counter

SOFFICE = shutil.which('soffice') or '/Applications/LibreOffice.app/Contents/MacOS/soffice'

docx_path = pathlib.Path(sys.argv[1])
prep_path = pathlib.Path(sys.argv[2])
orig_pdf  = pathlib.Path(sys.argv[3]) if len(sys.argv) > 3 else None

failures = []

def check(name, ok, detail=''):
    print(f'  [{"PASS" if ok else "FAIL"}] {name}' + (f' — {detail}' if detail else ''))
    if not ok:
        failures.append(name)

d = docx.Document(docx_path)
src = prep_path.read_text()

# ---- 1. code fidelity, line by line -----------------------------------------
blocks = re.findall(r'\\begin\{verbatim\}\n?(.*?)\\end\{verbatim\}', src, re.S)
src_lines = [l.rstrip() for b in blocks for l in b.split('\n') if l.strip()]
doc_lines = [l.rstrip() for p in d.paragraphs if p.style.name == 'Source Code'
             for l in p.text.split('\n') if l.strip()]
missing = [l for l in src_lines if l not in doc_lines]
check('code lines survive verbatim', not missing,
      f'{len(src_lines)} source vs {len(doc_lines)} docx' +
      (f'; first missing: {missing[0][:60]!r}' if missing else ''))

# ---- 2. element counts match the source -------------------------------------
n_tab_src = len(re.findall(r'\\begin\{table\}', src))
check('table count matches source', len(d.tables) == n_tab_src,
      f'{len(d.tables)} docx vs {n_tab_src} table envs (a HIGHER docx count '
      'usually means a dropped construct rendered as a phantom table)')
n_fig_src = len(re.findall(r'\\includegraphics', src))
check('image count matches source', len(d.inline_shapes) == n_fig_src,
      f'{len(d.inline_shapes)} docx vs {n_fig_src} includegraphics')

# ---- 3. code-in-figure trap --------------------------------------------------
bad = sum(1 for f in re.findall(r'\\begin\{figure\}.*?\\end\{figure\}', src, re.S)
          if '\\begin{verbatim}' in f)
check('no verbatim inside figure floats', bad == 0, f'{bad} figures hold code')

# ---- 4. heading hierarchy ----------------------------------------------------
c = Counter(p.style.name for p in d.paragraphs)
check('heading styles present', c.get('Heading 1', 0) + c.get('Heading 2', 0) > 0,
      f"H1={c.get('Heading 1',0)} H2={c.get('Heading 2',0)} H3={c.get('Heading 3',0)}")

# ---- 5. render, then text-level checks on the rendered PDF -------------------
out_pdf = docx_path.with_suffix('.pdf')
if not out_pdf.exists() or out_pdf.stat().st_mtime < docx_path.stat().st_mtime:
    subprocess.run([SOFFICE, '--headless', '--convert-to', 'pdf',
                    str(docx_path), '--outdir', str(docx_path.parent)],
                   capture_output=True)
txt = subprocess.run(['pdftotext', str(out_pdf), '-'],
                     capture_output=True, text=True).stdout

leaks = re.findall(r'\\[a-zA-Z]{2,}', txt)
check('LaTeX leakage low', len(leaks) <= 10,
      f'{len(leaks)} raw macros in rendered text (inspect each; code samples '
      'and Windows paths are legitimate)')
unresolved = re.findall(r'(?:Section|Chapter|Figure|Table|Appendix) \?', txt)
check('no unresolved references', not unresolved, f'{len(unresolved)} "?" refs')

# ---- 6. word retention vs the original PDF ----------------------------------
if orig_pdf and orig_pdf.exists():
    ot = subprocess.run(['pdftotext', str(orig_pdf), '-'],
                        capture_output=True, text=True).stdout
    # Filter to real words: raw tokens overcount the ORIGINAL, whose TOC dot
    # leaders, index entries, and page numbers do not (and should not) carry
    # into the docx. Measured this way the reference jobs landed at 94-97%.
    def words(s):
        return [w for w in re.sub(r'[^A-Za-z0-9_()#%.,=/-]+', ' ', s).split()
                if len(w) > 2]
    wn, wo = len(words(txt)), len(words(ot))
    pct = 100 * wn / max(wo, 1)
    check('word retention 90-105% of original', 90 <= pct <= 105,
          f'{pct:.1f}% ({wn} vs {wo}; shortfall should be running heads + TOC only)')

print()
if failures:
    print('FAILED:', ', '.join(failures)); sys.exit(1)
print('all checks passed')
